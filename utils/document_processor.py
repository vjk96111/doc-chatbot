"""
document_processor.py
Extracts text chunks and images from PDF and DOCX files.
"""

import io
import re
import base64
import zipfile
from typing import Dict, List


# ─────────────────────────────────── helpers ─────────────────────────────────

def _chunk_text(text: str, chunk_size: int = 400, overlap: int = 100) -> List[str]:
    """Split text into overlapping word-based chunks.
    Default overlap=100 words (≈ 500 chars) ensures ≥ 150-char overlap across all content.
    """
    words = text.split()
    if len(words) <= chunk_size:
        return [text]
    chunks: List[str] = []
    start = 0
    while start < len(words):
        end = min(start + chunk_size, len(words))
        chunks.append(" ".join(words[start:end]))
        if end >= len(words):
            break
        start += chunk_size - overlap
    return chunks


def _is_useful_image(img_bytes: bytes, ext: str) -> bool:
    """Filter out tiny icons / decorative images."""
    return len(img_bytes) >= 5_000 and ext.lower() in {
        "png", "jpeg", "jpg", "bmp", "tiff", "gif"
    }


def _ocr_page(page) -> str:
    """
    Render a PyMuPDF page as a greyscale image and extract text via Tesseract OCR.

    Called automatically when page.get_text() returns empty — handles scanned
    PDFs where the text layer is missing or uses a non-standard font encoding
    (e.g. Fujitsu PaperStream Capture, Adobe Acrobat scan-to-PDF, etc.).

    Returns empty string if pytesseract or Tesseract is not installed.
    """
    try:
        import pytesseract
        from PIL import Image as _PilImage
        import fitz as _fitz

        # Render at 150 DPI (scale factor ≈ 2.08 at PDF 72-DPI base).
        # Greyscale is faster for OCR and avoids colour-noise artefacts.
        mat = _fitz.Matrix(150 / 72, 150 / 72)
        pix = page.get_pixmap(matrix=mat, colorspace=_fitz.csGRAY)
        img = _PilImage.frombytes("L", [pix.width, pix.height], pix.samples)

        # --psm 3 = Fully automatic page segmentation (best for mixed layouts)
        return pytesseract.image_to_string(img, lang="eng", config="--psm 3").strip()

    except ImportError:
        # pytesseract not installed — caller handles empty return silently
        return ""
    except Exception:
        return ""


# ─────────────────────────────────── PDF ─────────────────────────────────────

def extract_from_pdf(file_bytes: bytes, doc_name: str = "") -> Dict:
    """
    Extract text chunks and images from a PDF.

    Returns:
        {
          "chunks": [{"text", "page", "source", "doc_name"}, ...],
          "images": {page_num: [{"data": b64_str, "ext": str, "page": int, "doc_name": str}, ...]}
        }
    """
    try:
        import fitz  # PyMuPDF
    except ImportError as exc:
        raise ImportError("PyMuPDF is required. Run: pip install PyMuPDF") from exc

    try:
        pdf = fitz.open(stream=file_bytes, filetype="pdf")
    except Exception as exc:
        raise ValueError(f"Could not open PDF '{doc_name}': {exc}") from exc

    all_chunks: List[Dict] = []
    images_by_page: Dict[int, List[Dict]] = {}
    ocr_pages: int = 0           # pages where OCR fallback was used

    for page_idx in range(len(pdf)):
        page = pdf[page_idx]
        page_num = page_idx + 1

        # ── text ──────────────────────────────────────────────────────────────
        raw_text = page.get_text("text").strip()

        # If PyMuPDF returns no text (scanned / non-standard font encoding),
        # fall back to Tesseract OCR rendered from the page image.
        if not raw_text:
            raw_text = _ocr_page(page)
            if raw_text:
                ocr_pages += 1

        if raw_text:
            for sub in _chunk_text(raw_text):
                all_chunks.append(
                    {
                        "text": sub,
                        "page": page_num,
                        "source": f"Page {page_num}",
                        "doc_name": doc_name,
                    }
                )

        # ── images ────────────────────────────────────────────────────────────
        page_imgs: List[Dict] = []
        for img_info in page.get_images(full=True):
            xref = img_info[0]
            try:
                base_img = pdf.extract_image(xref)
                img_bytes = base_img["image"]
                ext = base_img.get("ext", "png").lower().replace("jpg", "jpeg")
                if _is_useful_image(img_bytes, ext):
                    page_imgs.append(
                        {
                            "data": base64.b64encode(img_bytes).decode("utf-8"),
                            "ext": ext,
                            "page": page_num,
                            "doc_name": doc_name,
                        }
                    )
            except Exception:
                continue

        if page_imgs:
            images_by_page[page_num] = page_imgs

    pdf.close()
    return {"chunks": all_chunks, "images": images_by_page, "ocr_pages": ocr_pages}


# ─────────────────────────────────── DOCX ────────────────────────────────────

def _docx_rel_map(z: zipfile.ZipFile) -> Dict[str, str]:
    """Map relationship IDs to media paths inside the zip."""
    rel_file = "word/_rels/document.xml.rels"
    if rel_file not in z.namelist():
        return {}
    xml_text = z.read(rel_file).decode("utf-8", errors="replace")
    rel_map: Dict[str, str] = {}
    for m in re.finditer(r'Id="([^"]+)"[^>]*Target="([^"]+)"', xml_text):
        rid, target = m.group(1), m.group(2)
        if "media/" in target:
            # Normalise path: ../media/img.png → word/media/img.png
            clean = re.sub(r"^\.\./", "word/", target)
            rel_map[rid] = clean
    return rel_map


def _docx_image_order(z: zipfile.ZipFile) -> List[str]:
    """Return rIds of images in document order (order they appear in XML)."""
    doc_file = "word/document.xml"
    if doc_file not in z.namelist():
        return []
    xml_text = z.read(doc_file).decode("utf-8", errors="replace")
    return re.findall(r'r:embed="(rId\d+)"', xml_text)


def extract_from_docx(file_bytes: bytes, doc_name: str = "") -> Dict:
    """
    Extract text chunks and images from a DOCX file.

    Returns same structure as extract_from_pdf with "sections" instead of pages.
    """
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError("python-docx is required. Run: pip install python-docx") from exc

    doc = Document(io.BytesIO(file_bytes))

    # ── Build sections walking body in document order (paragraphs + tables) ──
    # doc.paragraphs silently skips table cells, so we walk the XML body directly.
    sections: List[Dict] = []       # [{"title", "text", "section_num"}]
    current_title = "Introduction"
    current_lines: List[str] = []
    section_num = 1

    _W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    def _flush():
        nonlocal section_num
        if current_lines:
            sections.append({
                "title":      current_title,
                "text":       "\n".join(current_lines),
                "section_num": section_num,
            })
            section_num += 1

    for child in doc.element.body:
        local = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if local == "p":
            # Collect all run text in paragraph order
            para_text = "".join(
                n.text or "" for n in child.iter(f"{{{_W}}}t")
            ).strip()
            if not para_text:
                continue
            # Check heading style
            pStyle = child.find(f".//{{{_W}}}pStyle")
            style_val = (pStyle.get(f"{{{_W}}}val") or "") if pStyle is not None else ""
            is_heading = style_val.lower().startswith("heading")
            if is_heading:
                _flush()
                current_title  = para_text
                current_lines  = []
            else:
                current_lines.append(para_text)

        elif local == "tbl":
            # Extract table rows as pipe-separated lines, injected into current section
            for tr in child.iter(f"{{{_W}}}tr"):
                cell_texts = []
                for tc in tr.iter(f"{{{_W}}}tc"):
                    cell_str = "".join(
                        n.text or "" for n in tc.iter(f"{{{_W}}}t")
                    ).strip()
                    if cell_str:
                        cell_texts.append(cell_str)
                if cell_texts:
                    current_lines.append(" | ".join(cell_texts))

    _flush()   # flush final section

    # Fallback: nothing extracted (unusual — all text was blank)
    if not sections:
        lines = []
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    lines.append(row_text)
        sections = [{"title": "Document", "text": "\n".join(lines), "section_num": 1}]

    total_sections = len(sections)

    # ── Build text chunks ─────────────────────────────────────────────────────
    all_chunks: List[Dict] = []
    for sec in sections:
        for sub in _chunk_text(sec["text"]):
            all_chunks.append(
                {
                    "text": sub,
                    "page": sec["section_num"],
                    "source": f"Section: {sec['title']}",
                    "doc_name": doc_name,
                }
            )

    # ── Extract images from the zip ───────────────────────────────────────────
    images_by_section: Dict[int, List[Dict]] = {}

    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
            rel_map = _docx_rel_map(z)
            ordered_rids = _docx_image_order(z)

            # Filter to image rIds only (in document order)
            img_rids_in_order = [
                rid for rid in ordered_rids if rid in rel_map
            ]

            total_imgs = len(img_rids_in_order)
            for idx, rid in enumerate(img_rids_in_order):
                media_path = rel_map[rid]
                if media_path not in z.namelist():
                    continue
                try:
                    img_bytes_raw = z.read(media_path)
                    ext = media_path.rsplit(".", 1)[-1].lower().replace("jpg", "jpeg")
                    if not _is_useful_image(img_bytes_raw, ext):
                        continue
                    b64 = base64.b64encode(img_bytes_raw).decode("utf-8")
                    # Distribute images proportionally across sections
                    sec_num = min(
                        int(idx * total_sections / max(total_imgs, 1)) + 1,
                        total_sections,
                    )
                    images_by_section.setdefault(sec_num, []).append(
                        {
                            "data": b64,
                            "ext": ext,
                            "page": sec_num,
                            "doc_name": doc_name,
                        }
                    )
                except Exception:
                    continue
    except Exception:
        pass  # Image extraction is best-effort

    return {"chunks": all_chunks, "images": images_by_section}
